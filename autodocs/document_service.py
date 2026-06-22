from __future__ import annotations

import re
import shutil
from datetime import date, datetime, time
from pathlib import Path
from typing import Iterable

from docx import Document
from openpyxl import Workbook, load_workbook
from openpyxl.styles import Font, PatternFill


PLACEHOLDER_PATTERN = re.compile(r"\[\[([^\[\]\r\n]+)\]\]")
INVALID_FILENAME_CHARS = re.compile(r'[<>:"/\\|?*\x00-\x1f]')
SYSTEM_PLACEHOLDERS = {"today", "current_day", "current_month", "current_year", "auto_number"}


def normalize_placeholder(name: str) -> str:
    return " ".join(name.strip().split())


def find_placeholders_in_text(text: str) -> list[str]:
    return [normalize_placeholder(match.group(1)) for match in PLACEHOLDER_PATTERN.finditer(text)]


def iter_document_paragraphs(document: Document) -> Iterable:
    for paragraph in document.paragraphs:
        yield paragraph

    for table in document.tables:
        yield from iter_table_paragraphs(table)

    for section in document.sections:
        for area in (section.header, section.footer):
            for paragraph in area.paragraphs:
                yield paragraph
            for table in area.tables:
                yield from iter_table_paragraphs(table)


def iter_table_paragraphs(table) -> Iterable:
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                yield paragraph
            for nested_table in cell.tables:
                yield from iter_table_paragraphs(nested_table)


def scan_docx_placeholders(path: str | Path) -> list[str]:
    document = Document(str(path))
    placeholders: set[str] = set()
    for paragraph in iter_document_paragraphs(document):
        placeholders.update(find_placeholders_in_text(paragraph.text))
    return sorted(placeholders, key=str.casefold)


def copy_template_to_library(source_path: str | Path, target_dir: str | Path, template_id: str) -> Path:
    target_dir = Path(target_dir)
    target_dir.mkdir(parents=True, exist_ok=True)
    target_path = target_dir / f"{template_id}.docx"
    shutil.copy2(source_path, target_path)
    return target_path


def render_filename(pattern: str, values: dict[str, object], fallback: str = "tai_lieu") -> str:
    def replace(match: re.Match) -> str:
        key = normalize_placeholder(match.group(1))
        value = values.get(key, "")
        return format_value_for_output(value).strip()

    rendered = PLACEHOLDER_PATTERN.sub(replace, pattern).strip()
    rendered = INVALID_FILENAME_CHARS.sub("_", rendered)
    rendered = re.sub(r"\s+", " ", rendered).strip(" .")
    return rendered or fallback


def format_value_for_output(value: object) -> str:
    if value is None:
        return ""
    if isinstance(value, datetime):
        if value.time().replace(microsecond=0) == time(0, 0):
            return value.strftime("%d/%m/%Y")
        return value.strftime("%d/%m/%Y %H:%M")
    if isinstance(value, date):
        return value.strftime("%d/%m/%Y")
    if isinstance(value, float) and value.is_integer():
        return str(int(value))
    return str(value)


def with_system_values(values: dict[str, object], auto_number: int | None = None) -> dict[str, object]:
    now = datetime.now()
    merged = dict(values)
    if not merged.get("today"):
        merged["today"] = now.strftime("%d/%m/%Y")
    if not merged.get("current_day"):
        merged["current_day"] = now.strftime("%d")
    if not merged.get("current_month"):
        merged["current_month"] = now.strftime("%m")
    if not merged.get("current_year"):
        merged["current_year"] = now.strftime("%Y")
    if auto_number is not None:
        if not merged.get("auto_number"):
            merged["auto_number"] = f"{auto_number:04d}"
    return merged


def replace_placeholders(template_path: str | Path, output_path: str | Path, values: dict[str, object]) -> None:
    document = Document(str(template_path))

    for paragraph in iter_document_paragraphs(document):
        replace_placeholders_in_paragraph(paragraph, values)

    output_path = Path(output_path)
    output_path.parent.mkdir(parents=True, exist_ok=True)
    document.save(str(output_path))


def replace_placeholders_in_paragraph(paragraph, values: dict[str, object]) -> None:
    text = paragraph.text
    if "[[" not in text or "]]" not in text:
        return

    matches = list(PLACEHOLDER_PATTERN.finditer(text))
    if not matches:
        return

    spans: list[tuple[int, int, object]] = []
    cursor = 0
    for run in paragraph.runs:
        start = cursor
        cursor += len(run.text)
        spans.append((start, cursor, run))

    def run_at(position: int) -> tuple[int, int]:
        for index, (start, end, _run) in enumerate(spans):
            if start <= position < end:
                return index, position - start
        if spans and position == spans[-1][1]:
            index = len(spans) - 1
            return index, spans[-1][1] - spans[-1][0]
        raise ValueError("Khong tim thay vi tri placeholder trong run.")

    for match in reversed(matches):
        key = normalize_placeholder(match.group(1))
        if key not in values:
            continue

        replacement = format_value_for_output(values[key])
        start_index, start_offset = run_at(match.start())
        end_index, end_offset = run_at(match.end())

        start_run = spans[start_index][2]
        end_run = spans[end_index][2]

        if start_index == end_index:
            current = start_run.text
            start_run.text = current[:start_offset] + replacement + current[end_offset:]
            continue

        start_run.text = start_run.text[:start_offset] + replacement
        end_run.text = end_run.text[end_offset:]
        for index in range(start_index + 1, end_index):
            spans[index][2].text = ""


def create_excel_sample(path: str | Path, placeholders: list[str]) -> None:
    workbook = Workbook()
    sheet = workbook.active
    sheet.title = "Du lieu"

    headers = placeholders or ["ten_truong"]
    for column, header in enumerate(headers, start=1):
        cell = sheet.cell(row=1, column=column, value=header)
        cell.font = Font(bold=True, color="FFFFFF")
        cell.fill = PatternFill("solid", fgColor="2563EB")
        sheet.column_dimensions[cell.column_letter].width = max(18, min(40, len(header) + 6))

    Path(path).parent.mkdir(parents=True, exist_ok=True)
    workbook.save(str(path))


def read_excel_rows(path: str | Path) -> list[dict[str, object]]:
    workbook = load_workbook(str(path), data_only=True)
    sheet = workbook.active
    rows = list(sheet.iter_rows(values_only=True))
    if not rows:
        return []

    headers = [normalize_placeholder(str(value)) if value is not None else "" for value in rows[0]]
    data_rows: list[dict[str, object]] = []
    for row in rows[1:]:
        item: dict[str, object] = {}
        has_value = False
        for index, header in enumerate(headers):
            if not header:
                continue
            value = row[index] if index < len(row) else None
            if value not in (None, ""):
                has_value = True
            item[header] = format_value_for_output(value)
        if has_value:
            data_rows.append(item)
    return data_rows


def convert_docx_to_pdf(docx_path: str | Path, pdf_path: str | Path) -> None:
    try:
        import win32com.client
    except ImportError as exc:
        raise RuntimeError("Chua cai pywin32 hoac khong co Microsoft Word de xuat PDF offline.") from exc

    word = win32com.client.Dispatch("Word.Application")
    word.Visible = False
    doc = None
    try:
        doc = word.Documents.Open(str(Path(docx_path).resolve()))
        doc.SaveAs(str(Path(pdf_path).resolve()), FileFormat=17)
    finally:
        if doc is not None:
            doc.Close(False)
        word.Quit()


def convert_pdf_to_images(pdf_path: str | Path, output_dir: str | Path, stem: str) -> list[Path]:
    try:
        import fitz
    except ImportError as exc:
        raise RuntimeError("Chua cai PyMuPDF de xuat anh tu PDF.") from exc

    output_dir = Path(output_dir)
    output_dir.mkdir(parents=True, exist_ok=True)
    document = fitz.open(str(pdf_path))
    image_paths: list[Path] = []
    try:
        for page_index in range(document.page_count):
            page = document.load_page(page_index)
            pixmap = page.get_pixmap(matrix=fitz.Matrix(2, 2), alpha=False)
            image_path = output_dir / f"{stem}_trang_{page_index + 1}.png"
            pixmap.save(str(image_path))
            image_paths.append(image_path)
    finally:
        document.close()
    return image_paths
