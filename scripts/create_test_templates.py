from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parent.parent
OUT = ROOT / "test_templates"

ACCENT = "2563EB"
INK = "172033"
MUTED = "667085"
HEADER_FILL = "EEF2F7"
LIGHT_FILL = "F8FAFC"
BORDER = "CBD5E1"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shading = tc_pr.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        tc_pr.append(shading)
    shading.set(qn("w:fill"), fill)


def set_cell_border(cell, color: str = BORDER) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = "w:{}".format(edge)
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "6")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_cell_margins(cell, top=90, start=120, bottom=90, end=120) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    margins = tc_pr.first_child_found_in("w:tcMar")
    if margins is None:
        margins = OxmlElement("w:tcMar")
        tc_pr.append(margins)
    for key, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        element = margins.find(qn(f"w:{key}"))
        if element is None:
            element = OxmlElement(f"w:{key}")
            margins.append(element)
        element.set(qn("w:w"), str(value))
        element.set(qn("w:type"), "dxa")


def set_table_width(table, widths_cm: list[float]) -> None:
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    for row in table.rows:
        for index, width in enumerate(widths_cm):
            if index < len(row.cells):
                row.cells[index].width = Cm(width)


def style_table(table, header_rows: int = 1, widths_cm: list[float] | None = None) -> None:
    if widths_cm:
        set_table_width(table, widths_cm)
    for row_index, row in enumerate(table.rows):
        for cell in row.cells:
            set_cell_border(cell)
            set_cell_margins(cell)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            if row_index < header_rows:
                set_cell_shading(cell, HEADER_FILL)
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.bold = True
                        run.font.color.rgb = RGBColor(23, 32, 51)
            elif row_index % 2 == 0:
                set_cell_shading(cell, LIGHT_FILL)


def set_doc_defaults(doc: Document, title: str, code: str) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)
    section.header_distance = Inches(0.2)
    section.footer_distance = Inches(0.2)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.1

    for name, size, color in (
        ("Heading 1", 16, ACCENT),
        ("Heading 2", 13, ACCENT),
        ("Heading 3", 12, "1F4D78"),
    ):
        style = styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(8)
        style.paragraph_format.space_after = Pt(5)

    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title_p.add_run(title)
    run.bold = True
    run.font.size = Pt(18)
    run.font.color.rgb = RGBColor.from_string(INK)


def add_meta_table(doc: Document, rows: list[tuple[str, str]]) -> None:
    table = doc.add_table(rows=len(rows), cols=2)
    for index, (label, value) in enumerate(rows):
        table.cell(index, 0).text = label
        table.cell(index, 1).text = value
    style_table(table, header_rows=0, widths_cm=[4.1, 11.9])
    for row in table.rows:
        set_cell_shading(row.cells[0], HEADER_FILL)
        row.cells[0].paragraphs[0].runs[0].bold = True


def add_signature_block(doc: Document, left: str, right: str) -> None:
    doc.add_paragraph()
    table = doc.add_table(rows=3, cols=2)
    table.cell(0, 0).text = left
    table.cell(0, 1).text = right
    table.cell(1, 0).text = "(Ky, ghi ro ho ten)"
    table.cell(1, 1).text = "(Ky, ghi ro ho ten)"
    table.cell(2, 0).text = "\n\n[[nguoi_ky_ben_a]]"
    table.cell(2, 1).text = "\n\n[[nguoi_ky_ben_b]]"
    style_table(table, header_rows=1, widths_cm=[8.0, 8.0])
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER


def build_sales_contract() -> None:
    doc = Document()
    set_doc_defaults(doc, "HOP DONG MUA BAN HANG HOA", "HDMB")
    add_meta_table(
        doc,
        [
            ("So hop dong", "[[so_hop_dong]]"),
            ("Ngay ky", "[[ngay_ky]]"),
            ("Ben ban", "[[ten_cong_ty_ban]] - MST: [[mst_ben_ban]]"),
            ("Ben mua", "[[ten_khach_hang]] - MST: [[mst_khach_hang]]"),
            ("Dia chi ben mua", "[[dia_chi_khach_hang]]"),
        ],
    )
    doc.add_heading("Dieu 1. Hang hoa va gia tri", level=1)
    table = doc.add_table(rows=4, cols=5)
    headers = ["STT", "Ten hang hoa", "So luong", "Don gia", "Thanh tien"]
    values = [
        ["1", "[[ten_hang_1]]", "[[so_luong_1]]", "[[don_gia_1]]", "[[thanh_tien_1]]"],
        ["2", "[[ten_hang_2]]", "[[so_luong_2]]", "[[don_gia_2]]", "[[thanh_tien_2]]"],
        ["Tong cong", "", "", "", "[[tong_gia_tri]]"],
    ]
    for c, header in enumerate(headers):
        table.cell(0, c).text = header
    for r, row in enumerate(values, start=1):
        for c, value in enumerate(row):
            table.cell(r, c).text = value
    style_table(table, widths_cm=[1.2, 6.0, 2.3, 3.0, 3.5])
    doc.add_heading("Dieu 2. Thanh toan va giao hang", level=1)
    doc.add_paragraph(
        "Ben mua thanh toan [[tong_gia_tri_bang_chu]] trong vong [[so_ngay_thanh_toan]] ngay "
        "ke tu ngay nhan du hoa don hop le. Dia diem giao hang: [[dia_diem_giao_hang]]."
    )
    doc.add_paragraph("Nguoi phu trach: [[nguoi_lien_he]] - Dien thoai: [[so_dien_thoai]].")
    add_signature_block(doc, "DAI DIEN BEN BAN", "DAI DIEN BEN MUA")
    doc.save(OUT / "01_hop_dong_mua_ban.docx")


def build_quote() -> None:
    doc = Document()
    set_doc_defaults(doc, "BAO GIA SAN PHAM / DICH VU", "BG")
    add_meta_table(
        doc,
        [
            ("Ma bao gia", "[[ma_bao_gia]]"),
            ("Khach hang", "[[ten_khach_hang]]"),
            ("Nguoi nhan", "[[nguoi_nhan_bao_gia]]"),
            ("Email", "[[email_khach_hang]]"),
            ("Hieu luc den", "[[ngay_het_han]]"),
        ],
    )
    doc.add_heading("Bang bao gia", level=1)
    table = doc.add_table(rows=5, cols=6)
    headers = ["Ma SP", "Hang muc", "Don vi", "SL", "Don gia", "Thanh tien"]
    rows = [
        ["[[ma_sp_1]]", "[[ten_san_pham_1]]", "[[don_vi_1]]", "[[so_luong_1]]", "[[don_gia_1]]", "[[thanh_tien_1]]"],
        ["[[ma_sp_2]]", "[[ten_san_pham_2]]", "[[don_vi_2]]", "[[so_luong_2]]", "[[don_gia_2]]", "[[thanh_tien_2]]"],
        ["[[ma_sp_3]]", "[[ten_san_pham_3]]", "[[don_vi_3]]", "[[so_luong_3]]", "[[don_gia_3]]", "[[thanh_tien_3]]"],
        ["", "Tong sau VAT [[vat_percent]]%", "", "", "", "[[tong_sau_vat]]"],
    ]
    for c, header in enumerate(headers):
        table.cell(0, c).text = header
    for r, row in enumerate(rows, start=1):
        for c, value in enumerate(row):
            table.cell(r, c).text = value
    style_table(table, widths_cm=[2.0, 5.4, 1.8, 1.2, 2.7, 2.9])
    doc.add_heading("Dieu kien ap dung", level=1)
    for text in [
        "Thoi gian giao hang du kien: [[thoi_gian_giao_hang]].",
        "Dieu kien thanh toan: [[dieu_kien_thanh_toan]].",
        "Ghi chu them: [[ghi_chu_bao_gia]].",
    ]:
        doc.add_paragraph(text, style="List Bullet")
    add_signature_block(doc, "NGUOI LAP BAO GIA", "KHACH HANG XAC NHAN")
    doc.save(OUT / "02_bao_gia_san_pham.docx")


def build_minutes() -> None:
    doc = Document()
    set_doc_defaults(doc, "BIEN BAN CUOC HOP", "BBH")
    add_meta_table(
        doc,
        [
            ("Chu de", "[[chu_de_cuoc_hop]]"),
            ("Thoi gian", "[[thoi_gian_hop]]"),
            ("Dia diem", "[[dia_diem_hop]]"),
            ("Chu tri", "[[nguoi_chu_tri]]"),
            ("Thu ky", "[[thu_ky]]"),
        ],
    )
    doc.add_heading("Thanh phan tham du", level=1)
    table = doc.add_table(rows=5, cols=4)
    headers = ["Ho ten", "Don vi", "Vai tro", "Trang thai"]
    rows = [
        ["[[nguoi_tham_du_1]]", "[[don_vi_1]]", "[[vai_tro_1]]", "[[trang_thai_1]]"],
        ["[[nguoi_tham_du_2]]", "[[don_vi_2]]", "[[vai_tro_2]]", "[[trang_thai_2]]"],
        ["[[nguoi_tham_du_3]]", "[[don_vi_3]]", "[[vai_tro_3]]", "[[trang_thai_3]]"],
        ["[[nguoi_tham_du_4]]", "[[don_vi_4]]", "[[vai_tro_4]]", "[[trang_thai_4]]"],
    ]
    for c, header in enumerate(headers):
        table.cell(0, c).text = header
    for r, row in enumerate(rows, start=1):
        for c, value in enumerate(row):
            table.cell(r, c).text = value
    style_table(table, widths_cm=[4.6, 4.1, 3.8, 3.5])
    doc.add_heading("Noi dung va ket luan", level=1)
    for label, value in [
        ("Noi dung chinh", "[[noi_dung_chinh]]"),
        ("Ket luan", "[[ket_luan_cuoc_hop]]"),
        ("Han hoan thanh", "[[han_hoan_thanh]]"),
    ]:
        p = doc.add_paragraph()
        p.add_run(f"{label}: ").bold = True
        p.add_run(value)
    doc.add_heading("Cong viec can theo doi", level=1)
    task_table = doc.add_table(rows=4, cols=4)
    headers = ["Cong viec", "Nguoi phu trach", "Han", "Trang thai"]
    rows = [
        ["[[cong_viec_1]]", "[[phu_trach_1]]", "[[deadline_1]]", "[[status_1]]"],
        ["[[cong_viec_2]]", "[[phu_trach_2]]", "[[deadline_2]]", "[[status_2]]"],
        ["[[cong_viec_3]]", "[[phu_trach_3]]", "[[deadline_3]]", "[[status_3]]"],
    ]
    for c, header in enumerate(headers):
        task_table.cell(0, c).text = header
    for r, row in enumerate(rows, start=1):
        for c, value in enumerate(row):
            task_table.cell(r, c).text = value
    style_table(task_table, widths_cm=[6.0, 4.0, 2.8, 3.2])
    add_signature_block(doc, "CHU TRI", "THU KY")
    doc.save(OUT / "03_bien_ban_cuoc_hop.docx")


def build_hr_decision() -> None:
    doc = Document()
    set_doc_defaults(doc, "QUYET DINH NHAN SU", "QDNS")
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run("So: [[so_quyet_dinh]]/[[current_year]]/QD-NS").bold = True
    doc.add_paragraph(
        "Can cu nhu cau nhan su cua [[ten_cong_ty]] va de xuat cua bo phan [[phong_ban_de_xuat]], "
        "Giam doc quyet dinh:"
    )
    doc.add_heading("Dieu 1. Thong tin nhan su", level=1)
    add_meta_table(
        doc,
        [
            ("Ho va ten", "[[ho_ten_nhan_su]]"),
            ("Ma nhan vien", "[[ma_nhan_vien]]"),
            ("Chuc danh moi", "[[chuc_danh_moi]]"),
            ("Phong ban", "[[phong_ban]]"),
            ("Ngay hieu luc", "[[ngay_hieu_luc]]"),
            ("Muc luong/Phu cap", "[[muc_luong]] / [[phu_cap]]"),
        ],
    )
    doc.add_heading("Dieu 2. Trach nhiem thi hanh", level=1)
    doc.add_paragraph(
        "[[ho_ten_nhan_su]], Truong bo phan [[phong_ban]], va cac phong ban lien quan co trach nhiem "
        "thi hanh quyet dinh nay ke tu ngay [[ngay_hieu_luc]]."
    )
    doc.add_paragraph("Noi nhan: [[noi_nhan]]")
    add_signature_block(doc, "NOI NHAN", "GIAM DOC")
    doc.save(OUT / "04_quyet_dinh_nhan_su.docx")


def build_invitation() -> None:
    doc = Document()
    set_doc_defaults(doc, "THU MOI THAM DU SU KIEN", "TM")
    add_meta_table(
        doc,
        [
            ("Nguoi nhan", "[[ten_nguoi_nhan]]"),
            ("Cong ty/Don vi", "[[don_vi_nguoi_nhan]]"),
            ("Su kien", "[[ten_su_kien]]"),
            ("Thoi gian", "[[thoi_gian_su_kien]]"),
            ("Dia diem", "[[dia_diem_su_kien]]"),
        ],
    )
    doc.add_paragraph(
        "Kinh gui [[ten_nguoi_nhan]],\n\n"
        "[[ten_cong_ty]] tran trong kinh moi Quy vi tham du [[ten_su_kien]]. "
        "Chuong trinh du kien bat dau luc [[gio_bat_dau]] va ket thuc luc [[gio_ket_thuc]]."
    )
    doc.add_heading("Noi dung chuong trinh", level=1)
    table = doc.add_table(rows=4, cols=3)
    headers = ["Thoi gian", "Noi dung", "Nguoi phu trach"]
    rows = [
        ["[[khung_gio_1]]", "[[noi_dung_1]]", "[[phu_trach_1]]"],
        ["[[khung_gio_2]]", "[[noi_dung_2]]", "[[phu_trach_2]]"],
        ["[[khung_gio_3]]", "[[noi_dung_3]]", "[[phu_trach_3]]"],
    ]
    for c, header in enumerate(headers):
        table.cell(0, c).text = header
    for r, row in enumerate(rows, start=1):
        for c, value in enumerate(row):
            table.cell(r, c).text = value
    style_table(table, widths_cm=[3.0, 8.2, 4.8])
    doc.add_paragraph("Vui long xac nhan tham du truoc ngay [[han_xac_nhan]] qua email [[email_lien_he]].")
    doc.add_paragraph("Tran trong,\n[[nguoi_gui_thu_moi]]\n[[chuc_danh_nguoi_gui]]")
    doc.save(OUT / "05_thu_moi_su_kien.docx")


def build_edge_cases() -> None:
    doc = Document()
    set_doc_defaults(doc, "MAU KIEM THU PLACEHOLDER DA DANG", "TEST")
    doc.add_paragraph("Placeholder lap lai: [[ten_khach_hang]] xuat hien lan 1.")
    doc.add_paragraph("Placeholder lap lai: [[ten_khach_hang]] xuat hien lan 2.")
    doc.add_paragraph("Bien he thong: [[today]], [[current_month]], [[current_year]], [[auto_number]].")
    doc.add_paragraph("Ten co dau gach duoi: [[ma_ho_so]], ten co khoang trang: [[ten nguoi dai dien]].")
    doc.add_heading("Bang long nhau don gian", level=1)
    table = doc.add_table(rows=3, cols=3)
    for c, header in enumerate(["Cot A", "Cot B", "Cot C"]):
        table.cell(0, c).text = header
    table.cell(1, 0).text = "[[field_a1]]"
    table.cell(1, 1).text = "[[field_b1]] va [[field_b1_phu]]"
    table.cell(1, 2).text = "[[field_c1]]"
    table.cell(2, 0).text = "[[field_a2]]"
    table.cell(2, 1).text = "[[field_b2]]"
    table.cell(2, 2).text = "[[field_c2]]"
    nested = table.cell(2, 2).add_table(rows=2, cols=2)
    nested.cell(0, 0).text = "Nested 1"
    nested.cell(0, 1).text = "[[nested_value_1]]"
    nested.cell(1, 0).text = "Nested 2"
    nested.cell(1, 1).text = "[[nested_value_2]]"
    style_table(table, widths_cm=[4.5, 5.5, 6.0])
    style_table(nested, widths_cm=[2.5, 3.0])
    doc.add_section(WD_SECTION_START.NEW_PAGE)
    doc.add_heading("Trang thu hai", level=1)
    doc.add_paragraph("Placeholder o trang moi: [[noi_dung_trang_2]].")
    doc.save(OUT / "06_mau_kiem_thu_da_dang.docx")


def main() -> None:
    OUT.mkdir(parents=True, exist_ok=True)
    for existing in OUT.glob("*.docx"):
        existing.unlink()
    build_sales_contract()
    build_quote()
    build_minutes()
    build_hr_decision()
    build_invitation()
    build_edge_cases()
    print(f"Created {len(list(OUT.glob('*.docx')))} files in {OUT}")


if __name__ == "__main__":
    main()
